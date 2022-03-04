def SalvarNoWord(TextoFormatado, Titulo, Imagens):
    try:
        DocumentoWord = Document()
    except:
        from docx import Document
        DocumentoWord = Document()

    try:
        LarguraDaFoto = Inches(3)
    except:
        from docx.shared import Inches
        LarguraDaFoto = Inches(3)

    DocumentoWord.add_heading(Titulo, 0)

    QuantidadeDeParagrafos = len(TextoFormatado)
    QuantidadeDeImagens = len(Imagens)

    proporcao = QuantidadeDeParagrafos // QuantidadeDeImagens

    indiceImagem = 0
    for i in range(len(TextoFormatado)):
        DocumentoWord.add_paragraph(TextoFormatado[i])
        if i > 0:
            if i % proporcao == 0:
                try:
                    DocumentoWord.add_picture(Imagens[indiceImagem], width=LarguraDaFoto)
                except:
                    pass
                indiceImagem += 1

    SalvarDocumentWord = r'C:\Users\Henrique\Desktop\trabalho\NOMEDOTRABALHO.docx'.replace('NOMEDOTRABALHO', Titulo)
    DocumentoWord.save(SalvarDocumentWord)